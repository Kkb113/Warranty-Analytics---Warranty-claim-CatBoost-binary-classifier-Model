"""Build the checked-in schema contract from the authoritative DOCX once.

This is a development-only extraction utility.  Runtime application code loads the
generated YAML and never parses the source document.  The command refuses to
overwrite an existing contract unless ``--force`` is supplied explicitly.
"""

from __future__ import annotations

import argparse
import hashlib
import re
from pathlib import Path
from typing import Any

import yaml
from docx import Document

CONTRACT_VERSION = "1.0.0"
DEFAULT_SOURCE = Path("warranty_analytics_schema_document.docx")
DEFAULT_OUTPUT = Path("contracts/warranty_analytics_schema_v1.yaml")
_TABLE_NAME_RE = re.compile(r"^(?P<schema>[^.]+)\.(?P<table>[^.]+)$")
_TYPE_RE = re.compile(r"^(?P<base>[A-Za-z0-9_]+)(?:\((?P<args>[^)]*)\))?$")


def _clean(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def _rows(table: Any) -> list[list[str]]:
    return [[_clean(cell.text) for cell in row.cells] for row in table.rows]


def _as_bool(value: str) -> bool:
    return _clean(value).casefold() in {"yes", "true", "1"}


def _as_int(value: str) -> int:
    return int(_clean(value).replace(",", ""))


def _optional(value: str) -> str | None:
    cleaned = _clean(value)
    return None if cleaned in {"", "-", "not documented"} else cleaned


def _parse_type(value: str) -> dict[str, Any]:
    match = _TYPE_RE.fullmatch(_clean(value))
    if match is None:
        raise ValueError(f"Unsupported SQL type in source document: {value!r}")
    base = match.group("base").lower()
    args = match.group("args")
    type_spec: dict[str, Any] = {
        "base_type": base,
        "max_length": None,
        "precision": None,
        "scale": None,
        "unicode": base.startswith("n"),
        "length_unit": None,
        "is_max": False,
    }
    if args is None:
        return type_spec
    parts = [part.strip() for part in args.split(",")]
    if base in {"nvarchar", "varchar", "nchar", "char", "varbinary", "binary"}:
        if parts[0].casefold() == "max":
            type_spec["is_max"] = True
        else:
            type_spec["max_length"] = int(parts[0])
        type_spec["length_unit"] = "characters" if base not in {"varbinary", "binary"} else "bytes"
    elif base in {"decimal", "numeric"}:
        if len(parts) != 2:
            raise ValueError(f"Decimal type must include precision and scale: {value!r}")
        type_spec["precision"] = int(parts[0])
        type_spec["scale"] = int(parts[1])
    else:
        raise ValueError(f"Unexpected SQL type arguments in source document: {value!r}")
    return type_spec


def _parse_columns(table: Any) -> list[dict[str, Any]]:
    rows = _rows(table)
    expected_header = [
        "Ord",
        "Column",
        "SQL type",
        "Null",
        "Identity",
        "Default / computed",
        "Collation",
    ]
    if rows[0] != expected_header:
        raise ValueError(f"Unexpected column table header: {rows[0]!r}")
    columns: list[dict[str, Any]] = []
    for row in rows[1:]:
        if len(row) != len(expected_header):
            raise ValueError(f"Unexpected column row: {row!r}")
        default_or_computed = _optional(row[5])
        columns.append(
            {
                "ordinal": _as_int(row[0]),
                "name": row[1],
                "sql_type": _parse_type(row[2]),
                "nullable": _as_bool(row[3]),
                "identity": _as_bool(row[4]),
                "computed": bool(
                    default_or_computed and "computed" in default_or_computed.casefold()
                ),
                "default": (
                    None
                    if default_or_computed is None or "computed" in default_or_computed.casefold()
                    else default_or_computed
                ),
                "collation": _optional(row[6]),
            }
        )
    return columns


def _parse_column_list(value: str) -> list[str]:
    cleaned = _clean(value).strip("[]")
    if not cleaned or cleaned == "-":
        return []
    return [part.strip().strip("[]") for part in cleaned.split(",")]


def _parse_key_columns(value: str) -> list[str]:
    columns: list[str] = []
    for part in _parse_column_list(value):
        columns.append(part.rsplit(" ", 1)[0] if " " in part else part)
    return columns


def _parse_foreign_keys(table: Any) -> list[dict[str, Any]]:
    rows = _rows(table)
    if rows[0] != [
        "Foreign key",
        "Referenced table",
        "Column mapping",
        "Delete",
        "Update",
        "Trusted",
    ]:
        raise ValueError(f"Unexpected foreign-key table header: {rows[0]!r}")
    foreign_keys: list[dict[str, Any]] = []
    for row in rows[1:]:
        mappings = [part.strip() for part in row[2].split(",")]
        parent_columns: list[str] = []
        referenced_columns: list[str] = []
        for mapping in mappings:
            sides = [side.strip() for side in mapping.split("->")]
            if len(sides) != 2:
                raise ValueError(f"Unexpected foreign-key mapping: {mapping!r}")
            parent_columns.append(sides[0])
            referenced_columns.append(sides[1])
        foreign_keys.append(
            {
                "name": row[0],
                "referenced_table": row[1],
                "parent_columns": parent_columns,
                "referenced_columns": referenced_columns,
                "on_delete": row[3].lower(),
                "on_update": row[4].lower(),
                "trusted": _as_bool(row[5]),
            }
        )
    return foreign_keys


def _parse_primary_key(table: Any) -> dict[str, Any]:
    rows = _rows(table)
    if rows[0] != ["Constraint", "Type", "Columns", "System named"]:
        raise ValueError(f"Unexpected constraint table header: {rows[0]!r}")
    primary_keys = [row for row in rows[1:] if row[1] == "PRIMARY_KEY_CONSTRAINT"]
    if len(primary_keys) != 1:
        raise ValueError("Each included table must have exactly one documented primary key.")
    row = primary_keys[0]
    return {
        "name": row[0],
        "columns": _parse_column_list(row[2]),
        "system_named": _as_bool(row[3]),
    }


def _parse_indexes(table: Any) -> list[dict[str, Any]]:
    rows = _rows(table)
    if rows[0] != ["Index", "Type", "Unique", "Key columns", "Included", "Disabled", "Filter"]:
        raise ValueError(f"Unexpected index table header: {rows[0]!r}")
    indexes: list[dict[str, Any]] = []
    for row in rows[1:]:
        indexes.append(
            {
                "name": row[0],
                "index_type": row[1].lower(),
                "unique": _as_bool(row[2]),
                "key_columns": _parse_key_columns(row[3]),
                "included_columns": _parse_column_list(row[4]),
                "disabled": _as_bool(row[5]),
                "filter_definition": _optional(row[6]),
            }
        )
    return indexes


def _find_table_indices(document: Any) -> list[int]:
    indices: list[int] = []
    for index, table in enumerate(document.tables[3:], start=3):
        rows = _rows(table)
        if rows and rows[0] == ["Table attribute", "Value"]:
            indices.append(index)
    return indices


def _table_section(document: Any, start: int, next_start: int | None) -> dict[str, Any]:
    end = next_start if next_start is not None else len(document.tables)
    section_tables = document.tables[start:end]
    attribute_rows = dict(_rows(section_tables[0])[1:])
    column_table = next(table for table in section_tables if _rows(table)[0][0] == "Ord")
    constraint_table = next(table for table in section_tables if _rows(table)[0][0] == "Constraint")
    index_table = next(table for table in section_tables if _rows(table)[0][0] == "Index")
    foreign_key_tables = [table for table in section_tables if _rows(table)[0][0] == "Foreign key"]
    storage = attribute_rows["Storage / temporal"]
    return {
        "attributes": attribute_rows,
        "columns": _parse_columns(column_table),
        "primary_key": _parse_primary_key(constraint_table),
        "foreign_keys": (_parse_foreign_keys(foreign_key_tables[0]) if foreign_key_tables else []),
        "indexes": _parse_indexes(index_table),
        "storage": storage,
        "memory_optimized": _as_bool(attribute_rows["Memory optimized"]),
        "file_table": _as_bool(attribute_rows["FileTable"]),
    }


def build_contract(root: Path, source: Path) -> dict[str, Any]:
    document = Document(str(source))
    control = dict(_rows(document.tables[0])[1:])
    overview_rows = _rows(document.tables[1])
    relationship_rows = _rows(document.tables[2])
    table_names = [row[0] for row in overview_rows[1:]]
    sections = _find_table_indices(document)
    if len(table_names) != len(sections):
        raise ValueError("The document table overview and table detail sections do not reconcile.")

    tables: list[dict[str, Any]] = []
    for position, fqn in enumerate(table_names):
        match = _TABLE_NAME_RE.fullmatch(fqn)
        if match is None:
            raise ValueError(f"Expected schema-qualified table name, got {fqn!r}")
        next_start = sections[position + 1] if position + 1 < len(sections) else None
        section = _table_section(document, sections[position], next_start)
        overview = overview_rows[position + 1]
        expected_rows = _as_int(overview[1])
        expected_columns = _as_int(overview[2])
        if expected_rows != _as_int(section["attributes"]["Current row-count estimate"]):
            raise ValueError(f"Row estimate mismatch for {fqn}.")
        if expected_columns != len(section["columns"]):
            raise ValueError(f"Column count mismatch for {fqn}.")
        section["fqn"] = fqn
        section["schema"] = match.group("schema")
        section["table"] = match.group("table")
        storage = section["storage"]
        section["estimated_rows"] = expected_rows
        section["column_count"] = expected_columns
        section["properties"] = {
            "storage_classification": section.pop("storage"),
            "temporal": storage.casefold() != "non_temporal_table",
            "memory_optimized": section.pop("memory_optimized"),
            "file_table": section.pop("file_table"),
        }
        tables.append(
            {
                "name": section.pop("fqn"),
                "schema": section.pop("schema"),
                "table": section.pop("table"),
                "estimated_rows": section.pop("estimated_rows"),
                "column_count": section.pop("column_count"),
                "primary_key": section.pop("primary_key"),
                "foreign_keys": section.pop("foreign_keys"),
                "indexes": section.pop("indexes"),
                "properties": section.pop("properties"),
                "columns": section.pop("columns"),
            }
        )

    source_relative = source.relative_to(root).as_posix()
    excluded_tables = [item.strip() for item in control["Excluded scope"].split(";")]
    relationships = relationship_rows[1:]
    total_rows = sum(table["estimated_rows"] for table in tables)
    total_columns = sum(table["column_count"] for table in tables)
    total_foreign_keys = sum(len(table["foreign_keys"]) for table in tables)
    if len(relationships) != total_foreign_keys:
        raise ValueError("The document relationship count and table foreign-key count differ.")

    return {
        "contract_version": CONTRACT_VERSION,
        "database": {
            "name": control["Database"],
            "platform": "sql_server",
        },
        "source": {
            "document": source_relative,
            "sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
            "extraction_date": control["Extracted"],
            "views_documented": 0,
            "sequences_documented": 0,
        },
        "summary": {
            "included_tables": len(tables),
            "included_columns": total_columns,
            "documented_foreign_keys": total_foreign_keys,
            "estimated_rows": total_rows,
        },
        "excluded_tables": excluded_tables,
        "tables": tables,
    }


def _resolve_path(root: Path, value: Path) -> Path:
    return value if value.is_absolute() else root / value


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument(
        "--force", action="store_true", help="Allow replacing an existing contract."
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="Rebuild and compare without writing the contract.",
    )
    arguments = parser.parse_args()
    root = Path(__file__).resolve().parents[1]
    source = _resolve_path(root, arguments.source).resolve()
    output = _resolve_path(root, arguments.output).resolve()
    if not source.is_file():
        parser.error(f"Source document does not exist: {source}")
    payload = build_contract(root, source)
    rendered = yaml.safe_dump(payload, sort_keys=False, allow_unicode=False)
    if arguments.check:
        if not output.is_file() or output.read_text(encoding="utf-8") != rendered:
            print(f"Contract is not reproducible at {output}.")
            return 1
        print(f"Contract is reproducible: {output}")
        return 0
    if output.exists() and not arguments.force:
        parser.error(f"Refusing to overwrite existing contract without --force: {output}")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(rendered, encoding="utf-8", newline="\n")
    print(f"Wrote schema contract: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
